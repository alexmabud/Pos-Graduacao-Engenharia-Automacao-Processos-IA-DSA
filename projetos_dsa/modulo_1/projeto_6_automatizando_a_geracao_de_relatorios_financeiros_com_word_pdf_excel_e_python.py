# Projeto 6 - Automatizando a Geracao de Relatorios Financeiros com Word, PDF, Excel e Python
# Script Unificado: Gera multiplos relatorios em Word e converte PDF/DOCX para Excel

# Imports para geracao de relatorios Word
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Imports para conversao PDF/DOCX para Excel
import os
import pdfplumber
import pandas as pd
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk

# Caminho onde os documentos serao salvos
CAMINHO_SAIDA = r"C:\Users\User\OneDrive\Documentos\Python\Dev_Python\Abud Python Learning\DSA\Módulo_1-Automação-Excel-e-Engenharia-de-Dados\9_automacao_de_processos_com_excel_em_python"


# ============================================================================
# PARTE 1: FUNCOES PARA GERACAO DE RELATORIOS WORD
# ============================================================================

def criar_tabela_formatada(doc, titulo, dados, adicionar_quebra=False):
    """
    Cria uma tabela formatada no documento Word.

    Args:
        doc: Objeto Document do python-docx
        titulo: Título da seção (opcional)
        dados: Lista de listas com os dados da tabela
        adicionar_quebra: Se True, adiciona quebra de página após a tabela
    """

    # Adiciona título se fornecido
    if titulo:
        doc.add_heading(titulo, level=2)

    # Define cabeçalhos
    cabecalhos = ['Data', 'Descrição', 'Receitas', 'Despesas', 'Saldo Anterior', 'Saldo Atual']

    # Cria tabela (1 linha cabeçalho + N linhas de dados)
    table = doc.add_table(rows=len(dados) + 1, cols=len(cabecalhos))

    # Acessa propriedades da tabela para formatação
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    # Define bordas XML
    tbl_borders = parse_xml(
        r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/>'
        r'<w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/>'
        r'<w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>'
        % nsdecls('w')
    )
    tbl_pr.append(tbl_borders)

    # Preenche cabeçalhos
    header_cells = table.rows[0].cells
    for i, cabecalho in enumerate(cabecalhos):
        header_cells[i].text = cabecalho
        # Formata tamanho de fonte do cabeçalho
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Preenche dados
    for row_idx, row_data in enumerate(dados, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)

    # Adiciona quebra de página se solicitado
    if adicionar_quebra:
        doc.add_page_break()


def gerar_relatorio_simples(nome_arquivo="documento2.docx", caminho=None):
    """
    Gera relatorio simples com UMA tabela (10 registros).
    Equivalente ao gera_documento2.py

    Args:
        nome_arquivo: Nome do arquivo a salvar
        caminho: Caminho para salvar (usa CAMINHO_SAIDA se None)
    """
    if caminho is None:
        caminho = CAMINHO_SAIDA

    caminho_completo = os.path.join(caminho, nome_arquivo)

    doc = Document()
    doc.add_heading('Relatorio Financeiro', level=1)

    # Dados do período 01-10 de agosto
    dados = [
        ['01/08/2025', 'Venda Produto A', '5000', '', '10000', '15000'],
        ['02/08/2025', 'Compra Insumos', '', '2000', '15000', '13000'],
        ['03/08/2025', 'Venda Produto B', '3000', '', '13000', '16000'],
        ['04/08/2025', 'Despesas Operacionais', '', '1500', '16000', '14500'],
        ['05/08/2025', 'Venda Produto C', '7000', '', '14500', '21500'],
        ['06/08/2025', 'Despesas de Marketing', '', '2500', '21500', '19000'],
        ['07/08/2025', 'Venda Produto D', '4000', '', '19000', '23000'],
        ['08/08/2025', 'Compra de Equipamentos', '', '6000', '23000', '17000'],
        ['09/08/2025', 'Venda Produto E', '3500', '', '17000', '20500'],
        ['10/08/2024', 'Despesas de Transporte', '', '1200', '20500', '19300']
    ]

    criar_tabela_formatada(doc, None, dados, adicionar_quebra=False)
    doc.save(caminho_completo)

    return f"[OK] Relatorio simples criado: {caminho_completo}"


def gerar_relatorio_completo(nome_arquivo="documento3.docx", caminho=None):
    """
    Gera relatorio completo com DUAS tabelas (20 registros em 2 paginas).
    Equivalente ao gera_documento3.py

    Args:
        nome_arquivo: Nome do arquivo a salvar
        caminho: Caminho para salvar (usa CAMINHO_SAIDA se None)
    """
    if caminho is None:
        caminho = CAMINHO_SAIDA

    caminho_completo = os.path.join(caminho, nome_arquivo)

    doc = Document()
    doc.add_heading('Relatorio Financeiro', level=1)

    # Dados do período 01-10 de agosto (Primeira tabela)
    dados_parte1 = [
        ['01/08/2025', 'Venda Produto A', '5000', '', '10000', '15000'],
        ['02/08/2025', 'Compra Insumos', '', '2000', '15000', '13000'],
        ['03/08/2025', 'Venda Produto B', '3000', '', '13000', '16000'],
        ['04/08/2025', 'Despesas Operacionais', '', '1500', '16000', '14500'],
        ['05/08/2025', 'Venda Produto C', '7000', '', '14500', '21500'],
        ['06/08/2025', 'Despesas de Marketing', '', '2500', '21500', '19000'],
        ['07/08/2025', 'Venda Produto D', '4000', '', '19000', '23000'],
        ['08/08/2025', 'Compra de Equipamentos', '', '6000', '23000', '17000'],
        ['09/08/2025', 'Venda Produto E', '3500', '', '17000', '20500'],
        ['10/08/2024', 'Despesas de Transporte', '', '1200', '20500', '19300']
    ]

    # Primeira tabela
    criar_tabela_formatada(doc, "Período: 01-10 de Agosto de 2025", dados_parte1, adicionar_quebra=True)

    # Dados do período 11-20 de agosto (Segunda tabela)
    dados_parte2 = [
        ['11/08/2025', 'Venda Produto F', '8000', '', '19300', '27300'],
        ['12/08/2025', 'Compra de Matéria-Prima', '', '4000', '27300', '23300'],
        ['13/08/2025', 'Venda Produto G', '4500', '', '23300', '27800'],
        ['14/08/2025', 'Despesas de Manutenção', '', '1700', '27800', '26100'],
        ['15/08/2025', 'Venda Produto H', '6000', '', '26100', '32100'],
        ['16/08/2025', 'Despesas Diversas', '', '2300', '32100', '29800'],
        ['17/08/2025', 'Venda Produto I', '7500', '', '29800', '37300'],
        ['18/08/2025', 'Compra de Ferramentas', '', '3200', '37300', '34100'],
        ['19/08/2025', 'Venda Produto J', '5000', '', '34100', '39100'],
        ['20/08/2025', 'Despesas Gerais', '', '1800', '39100', '37300']
    ]

    # Segunda tabela
    criar_tabela_formatada(doc, "Periodo: 11-20 de Agosto de 2025", dados_parte2, adicionar_quebra=False)

    doc.save(caminho_completo)

    return f"[OK] Relatorio completo criado: {caminho_completo}"


def gerar_relatorio_personalizado(nome_arquivo="documento_custom.docx", num_tabelas=2, caminho=None):
    """
    Gera relatorio personalizado com numero flexivel de tabelas.

    Args:
        nome_arquivo: Nome do arquivo a salvar
        num_tabelas: Numero de tabelas a gerar (1 ou 2)
        caminho: Caminho para salvar (usa CAMINHO_SAIDA se None)
    """
    if caminho is None:
        caminho = CAMINHO_SAIDA

    if num_tabelas == 1:
        return gerar_relatorio_simples(nome_arquivo, caminho)
    elif num_tabelas == 2:
        return gerar_relatorio_completo(nome_arquivo, caminho)
    else:
        return "Erro: num_tabelas deve ser 1 ou 2"


# ============================================================================
# PARTE 2: FUNCOES PARA CONVERSAO PDF/DOCX PARA EXCEL
# ============================================================================

def dsa_padroniza_header(header):
    """Padroniza os cabecalhos das tabelas."""
    if not header:
        return []
    return [col.strip().lower().replace("\n", " ") if col else '' for col in header]


def dsa_verifica_coluna_unica(columns):
    """Verifica e renomeia colunas duplicadas em um DataFrame."""
    seen = {}
    for i, col in enumerate(columns):
        if col in seen:
            seen[col] += 1
            columns[i] = f"{col}_{seen[col]}"
        else:
            seen[col] = 0
    return columns


def dsa_limpa_alinha_tabela(df, combined_df):
    """Alinha colunas de um DataFrame com um DataFrame combinado."""
    df.columns = dsa_verifica_coluna_unica(list(df.columns))

    for col in combined_df.columns:
        if col not in df.columns:
            df[col] = ''

    df = df[combined_df.columns]
    return df


def dsa_remove_linhas_branco(df):
    """Remove linhas em branco de um DataFrame."""
    df = df.dropna(how='all')
    df = df[df.iloc[:, 0].notna() & (df.iloc[:, 0] != '')]
    return df


def dsa_converte_pdf_excel(pdf_path, excel_path):
    """Converte tabelas de PDF para Excel."""
    print(f"Pasta do PDF: {pdf_path}")
    print(f"Pasta do Excel: {excel_path}")

    if not os.access(os.path.dirname(excel_path), os.W_OK):
        raise PermissionError(f"Nao e possivel gravar na pasta: {os.path.dirname(excel_path)}")

    tables_by_header = {}

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()

            for table in tables:
                if table and len(table) > 1 and table[0]:
                    header = dsa_padroniza_header(table[0])
                    header_tuple = tuple(header)

                    if any(header):
                        df = pd.DataFrame(table[1:], columns=header)
                        df.columns = dsa_verifica_coluna_unica(list(df.columns))
                        print(f"Pagina {page_num + 1} titulos das colunas: {header}")

                        if header_tuple in tables_by_header:
                            df = dsa_limpa_alinha_tabela(df, tables_by_header[header_tuple])
                            tables_by_header[header_tuple] = pd.concat([tables_by_header[header_tuple], df], ignore_index=True)
                        else:
                            tables_by_header[header_tuple] = df
                    else:
                        print(f"Pular uma tabela na pagina {page_num + 1} devido a cabecalho invalido ou vazio.")
                else:
                    print(f"Pular uma tabela na pagina {page_num + 1} devido a dados ausentes ou invalidos.")

    if tables_by_header:
        try:
            with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
                for i, (header_tuple, combined_table) in enumerate(tables_by_header.items()):
                    combined_table.columns = dsa_verifica_coluna_unica(list(combined_table.columns))
                    combined_table = dsa_remove_linhas_branco(combined_table)
                    sheet_name = f"Tabela_{i+1}"
                    combined_table.to_excel(writer, sheet_name=sheet_name, index=False)
            messagebox.showinfo("Sucesso", f"As tabelas foram extraidas e salvas com sucesso em {excel_path}")
        except Exception as e:
            print(f"Falha ao salvar o arquivo Excel: {e}")
            messagebox.showerror("Erro", f"Falha ao salvar o arquivo Excel: {e}")
    else:
        messagebox.showinfo("Nenhuma Tabela Encontrada", "Nao ha tabela no arquivo PDF.")


def dsa_converte_docx_excel(docx_path, excel_path):
    """Converte tabelas de DOCX para Excel."""
    print(f"Pasta do DOCX: {docx_path}")
    print(f"Pasta do Excel: {excel_path}")

    if not os.access(os.path.dirname(excel_path), os.W_OK):
        raise PermissionError(f"Nao e possivel gravar na pasta: {os.path.dirname(excel_path)}")

    doc = Document(docx_path)
    tables_by_header = {}

    for table_num, table in enumerate(doc.tables):
        data = [[cell.text for cell in row.cells] for row in table.rows]

        if data and len(data) > 1 and data[0]:
            header = dsa_padroniza_header(data[0])
            header_tuple = tuple(header)

            if any(header):
                df = pd.DataFrame(data[1:], columns=header)
                df.columns = dsa_verifica_coluna_unica(list(df.columns))
                print(f"Tabela {table_num + 1} titulos das colunas: {header}")

                if header_tuple in tables_by_header:
                    df = dsa_limpa_alinha_tabela(df, tables_by_header[header_tuple])
                    tables_by_header[header_tuple] = pd.concat([tables_by_header[header_tuple], df], ignore_index=True)
                else:
                    tables_by_header[header_tuple] = df
            else:
                print(f"Pular uma tabela devido a cabecalho invalido ou vazio.")
        else:
            print(f"Pular uma tabela devido a dados ausentes ou invalidos.")

    if tables_by_header:
        try:
            with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
                for i, (header_tuple, combined_table) in enumerate(tables_by_header.items()):
                    combined_table.columns = dsa_verifica_coluna_unica(list(combined_table.columns))
                    combined_table = dsa_remove_linhas_branco(combined_table)
                    sheet_name = f"Tabela_{i+1}"
                    combined_table.to_excel(writer, sheet_name=sheet_name, index=False)
            messagebox.showinfo("Sucesso", f"As tabelas foram extraidas e salvas com sucesso em {excel_path}")
        except Exception as e:
            print(f"Falha ao salvar o arquivo Excel: {e}")
            messagebox.showerror("Erro", f"Falha ao salvar o arquivo Excel: {e}")
    else:
        messagebox.showinfo("Nenhuma Tabela Encontrada", "Nao ha tabela no arquivo DOCX.")


# ============================================================================
# PARTE 3: INTERFACE GRAFICA PARA CONVERSAO
# ============================================================================

def dsa_seleciona_arquivo():
    """Seleciona o arquivo PDF ou DOCX."""
    file_path = filedialog.askopenfilename(title="Selecione o Arquivo", filetypes=[("PDF Files", "*.pdf"), ("Word Files", "*.docx")])

    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)


def dsa_seleciona_local_salvar_excel():
    """Seleciona o local para salvar o arquivo Excel."""
    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", title="Salvar Arquivo Excel", filetypes=[("Excel Files", "*.xlsx")])

    if file_path:
        excel_entry.delete(0, tk.END)
        excel_entry.insert(0, file_path)


def dsa_inicia_conversao():
    """Inicia o processo de conversao com base no tipo de arquivo selecionado."""
    file_path = file_entry.get()
    excel_path = excel_entry.get()

    if not file_path or not excel_path:
        messagebox.showwarning("Entrada necessaria", "Selecione o arquivo e o destino para salvar o arquivo Excel.")
    else:
        if file_path.endswith(".pdf"):
            dsa_converte_pdf_excel(file_path, excel_path)
        elif file_path.endswith(".docx"):
            dsa_converte_docx_excel(file_path, excel_path)
        else:
            messagebox.showerror("Erro", "Tipo de arquivo nao suportado. Selecione um arquivo PDF ou DOCX.")


def abrir_interface_conversao():
    """Abre interface grafica para conversao PDF/DOCX para Excel."""
    global file_entry, excel_entry

    root = tk.Tk()
    root.title("DSA - Projeto 6 - Conversao para Excel")

    tk.Label(root, text="Selecione o Arquivo (PDF ou DOCX):").grid(row=0, column=0, padx=10, pady=10)
    file_entry = tk.Entry(root, width=50)
    file_entry.grid(row=0, column=1, padx=10, pady=10)

    file_button = ttk.Button(root, text="Browse", command=dsa_seleciona_arquivo)
    file_button.grid(row=0, column=2, padx=10, pady=10)

    tk.Label(root, text="Selecione o Destino Para Salvar o Excel:").grid(row=1, column=0, padx=10, pady=10)
    excel_entry = tk.Entry(root, width=50)
    excel_entry.grid(row=1, column=1, padx=10, pady=10)

    excel_button = ttk.Button(root, text="Browse", command=dsa_seleciona_local_salvar_excel)
    excel_button.grid(row=1, column=2, padx=10, pady=10)

    convert_button = ttk.Button(root, text="Extrair Tabela", command=dsa_inicia_conversao)
    convert_button.grid(row=2, columnspan=3, pady=20)

    root.mainloop()


# ============================================================================
# EXECUCAO PRINCIPAL
# ============================================================================

if __name__ == "__main__":
    print("=" * 70)
    print("PROJETO 6 - AUTOMACAO DE RELATORIOS FINANCEIROS")
    print("=" * 70)
    print(f"\nCaminho de saida: {CAMINHO_SAIDA}\n")
    print("Escolha uma opcao:")
    print("[1] Gerar relatorios Word automaticamente")
    print("[2] Converter PDF/DOCX para Excel (Interface Grafica)")
    print("[0] Sair")
    print("=" * 70)

    opcao = input("\nDigite sua opcao: ").strip()

    if opcao == "1":
        print("\n" + "=" * 70)
        print("GERACAO DE RELATORIOS WORD")
        print("=" * 70)

        # Opcao 1: Gerar relatorio simples (1 tabela)
        print("\n[1] Gerando relatorio simples (1 tabela, 10 registros)...")
        resultado1 = gerar_relatorio_simples("documento2.docx", CAMINHO_SAIDA)
        print(resultado1)

        # Opcao 2: Gerar relatorio completo (2 tabelas)
        print("\n[2] Gerando relatorio completo (2 tabelas, 20 registros)...")
        resultado2 = gerar_relatorio_completo("documento3.docx", CAMINHO_SAIDA)
        print(resultado2)

        print("\n" + "=" * 70)
        print("RESUMO:")
        print("=" * 70)
        print(resultado1)
        print(resultado2)
        print("\nAmbos os relatorios foram gerados com sucesso!")
        print("=" * 70)

    elif opcao == "2":
        print("\n[INFO] Abrindo interface grafica para conversao...")
        print("[INFO] Selecione arquivo PDF ou DOCX e escolha onde salvar o Excel")
        abrir_interface_conversao()

    elif opcao == "0":
        print("\n[INFO] Saindo do programa...")

    else:
        print("\n[ERRO] Opcao invalida! Execute o programa novamente.")
