# Projeto 5 - Detecção Automática de Anomalias em Logs de Pipelines de Dados
# Script Unificado (Full Stack ETL + IA)

# Imports
import pandas as pd
import os
from docx import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
# Usando a biblioteca moderna que você já tinha (langchain_ollama é mais atual que langchain_community)
from langchain_ollama import OllamaLLM 

# ==============================
# CONFIGURAÇÃO GERAL E CAMINHOS
# ==============================

# Caminho Base (Centralizado para facilitar manutenção)
DIRETORIO_BASE = r"C:\Users\User\OneDrive\Documentos\Python\Dev_Python\Abud Python Learning\DSA\Módulo_1-Automação-Excel-e-Engenharia-de-Dados\8_automatizando_a_engenharia_de_dados_com_ia"

# Arquivos de Entrada e Saída
ARQUIVO_LOGS = os.path.join(DIRETORIO_BASE, "logs_pipeline.xlsx")
ARQUIVO_RELATORIO = os.path.join(DIRETORIO_BASE, "projeto5-resultado.docx")

# Configuração do Modelo LLM (Llama3 via Ollama)
MODELO_IA = "llama3"

# ==============================
# FUNÇÃO 1: VALIDAÇÃO DOS DADOS
# ==============================
def dsa_valida_arquivo(caminho_arquivo):
    """
    Verifica se o arquivo Excel existe, é legível e contém dados.
    """
    print(f"--- Etapa 1: Validando arquivo {os.path.basename(caminho_arquivo)} ---")
    
    if not os.path.exists(caminho_arquivo):
        return False, "Erro: O arquivo não foi encontrado no diretório especificado."

    try:
        df = pd.read_excel(caminho_arquivo)
        
        if df.empty:
            return False, "Erro: O arquivo existe, mas está vazio."
        
        num_rows = df.shape[0]
        print(f"Arquivo validado com sucesso! Contém {num_rows} registros para análise.")
        return True, df # Retorna True e o próprio DataFrame carregado
    
    except Exception as e:
        return False, f"Erro crítico ao ler o arquivo: {str(e)}"

# ==============================
# FUNÇÃO 2: ANÁLISE COM IA (LLM)
# ==============================
def dsa_analisa_pipeline(df_dados):
    """
    Processa cada linha do log usando Llama3 e gera um relatório Word.
    """
    print(f"\n--- Etapa 2: Iniciando Análise Inteligente com {MODELO_IA} ---")

    # 1. Configuração da IA
    llm = OllamaLLM(model=MODELO_IA)
    output_parser = StrOutputParser()
    
    # 2. Criação do Template de Prompt (Engenharia de Prompt)
    prompt = ChatPromptTemplate.from_messages([
        ("system", "Você é um engenheiro de dados sênior especializado em observabilidade. Analise o log de execução e identifique anomalias ou confirme o sucesso em português do Brasil."),
        ("user", "Dados do Log: {question}")
    ])

    # 3. Pipeline LangChain (Prompt -> LLM -> Parser)
    chain = prompt | llm | output_parser

    # 4. Preparação do Documento Word
    document = Document()
    document.add_heading('Relatório de Análise de Anomalias em Pipelines', 0)

    resultados = []
    total_linhas = len(df_dados)

    # 5. Iteração (Loop) pelos dados
    for index, row in df_dados.iterrows():
        print(f"Processando registro {index + 1}/{total_linhas} (Execution ID: {row['Execution_ID']})...")

        # Extração de dados (Mapeamento das colunas do Excel)
        # Nota: Certifique-se que seu Excel tem exatamente estas colunas nesta ordem
        pipeline_id = row['Pipeline_ID']
        execution_id = row['Execution_ID']
        status = row['Status']
        tempo_execucao = row['Execution_Time_Minutes']
        sistema_op = row['Operating_System']
        tipo_proc = row['Processing_Type']
        tentativa = row['Attempt_Number']

        # Montagem da consulta para a IA
        consulta = (
            f"Analise esta execução: Pipeline ID: {pipeline_id}, Execution ID: {execution_id}, "
            f"Status: {status}, Tempo: {tempo_execucao} min, OS: {sistema_op}, "
            f"Tipo: {tipo_proc}, Tentativa: {tentativa}. "
            f"Há algo anômalo ou está tudo certo?"
        )

        # Invocação da IA
        resposta_ia = chain.invoke({'question': consulta})
        
        # Salva na lista e no documento
        resultados.append(resposta_ia)
        
        # Adiciona no Word: Título da execução + Resposta da IA
        document.add_heading(f"Análise da Execução {execution_id}", level=1)
        document.add_paragraph(resposta_ia)
        document.add_paragraph("-" * 50) # Separador visual

    # 6. Salvar Relatório Final
    try:
        document.save(ARQUIVO_RELATORIO)
        print(f"\n--- Etapa 3: Relatório salvo com sucesso em: ---")
        print(ARQUIVO_RELATORIO)
    except Exception as e:
        print(f"Erro ao salvar o documento Word: {e}")

    return resultados

# ==============================
# EXECUÇÃO PRINCIPAL (MAIN)
# ==============================
if __name__ == "__main__":
    
    print(">>> INICIANDO SISTEMA DE DETECÇÃO DE ANOMALIAS <<<\n")

    # Passo 1: Validar Arquivo
    sucesso, resultado_validacao = dsa_valida_arquivo(ARQUIVO_LOGS)

    if sucesso:
        # Se o arquivo for válido, 'resultado_validacao' contém o DataFrame
        df_pipeline = resultado_validacao
        
        # Passo 2: Executar Análise
        dsa_analisa_pipeline(df_pipeline)
        
        print("\n>>> PROCESSO CONCLUÍDO COM SUCESSO <<<")
    else:
        # Se falhar, exibe o erro e para
        print(f"\nFalha na validação: {resultado_validacao}")