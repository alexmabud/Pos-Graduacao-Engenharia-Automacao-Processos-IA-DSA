# Projeto 6 - Automatizando a Geracao de Relatorios Financeiros com Word, PDF, Excel e Python
# Script Unificado: Gera multiplos relatorios em Word com tabelas financeiras

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# Caminho onde os documentos serao salvos
CAMINHO_SAIDA = r"C:\Users\User\OneDrive\Documentos\Python\Dev_Python\Abud Python Learning\DSA\Módulo_1-Automação-Excel-e-Engenharia-de-Dados\9_automacao_de_processos_com_excel_em_python"


def criar_tabela_formatada(doc, titulo, dados, adicionar_quebra=False):
    """
    Cria uma tabela formatada no documento Word.

    Args:
        doc: Objeto Document do python-docx
        titulo: Título da seção (opcional)
        dados: Lista de listas com os dados da tabela
        adicionar_quebra: Se True, adiciona quebra de página após a tabela
    """

    # Adiciona título se fornecido
    if titulo:
        doc.add_heading(titulo, level=2)

    # Define cabeçalhos
    cabecalhos = ['Data', 'Descrição', 'Receitas', 'Despesas', 'Saldo Anterior', 'Saldo Atual']

    # Cria tabela (1 linha cabeçalho + N linhas de dados)
    table = doc.add_table(rows=len(dados) + 1, cols=len(cabecalhos))

    # Acessa propriedades da tabela para formatação
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    # Define bordas XML
    tbl_borders = parse_xml(
        r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/>'
        r'<w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/>'
        r'<w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>'
        % nsdecls('w')
    )
    tbl_pr.append(tbl_borders)

    # Preenche cabeçalhos
    header_cells = table.rows[0].cells
    for i, cabecalho in enumerate(cabecalhos):
        header_cells[i].text = cabecalho
        # Formata tamanho de fonte do cabeçalho
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Preenche dados
    for row_idx, row_data in enumerate(dados, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)

    # Adiciona quebra de página se solicitado
    if adicionar_quebra:
        doc.add_page_break()


def gerar_relatorio_simples(nome_arquivo="documento2.docx", caminho=None):
    """
    Gera relatorio simples com UMA tabela (10 registros).
    Equivalente ao gera_documento2.py

    Args:
        nome_arquivo: Nome do arquivo a salvar
        caminho: Caminho para salvar (usa CAMINHO_SAIDA se None)
    """
    if caminho is None:
        caminho = CAMINHO_SAIDA

    caminho_completo = os.path.join(caminho, nome_arquivo)

    doc = Document()
    doc.add_heading('Relatorio Financeiro', level=1)

    # Dados do período 01-10 de agosto
    dados = [
        ['01/08/2025', 'Venda Produto A', '5000', '', '10000', '15000'],
        ['02/08/2025', 'Compra Insumos', '', '2000', '15000', '13000'],
        ['03/08/2025', 'Venda Produto B', '3000', '', '13000', '16000'],
        ['04/08/2025', 'Despesas Operacionais', '', '1500', '16000', '14500'],
        ['05/08/2025', 'Venda Produto C', '7000', '', '14500', '21500'],
        ['06/08/2025', 'Despesas de Marketing', '', '2500', '21500', '19000'],
        ['07/08/2025', 'Venda Produto D', '4000', '', '19000', '23000'],
        ['08/08/2025', 'Compra de Equipamentos', '', '6000', '23000', '17000'],
        ['09/08/2025', 'Venda Produto E', '3500', '', '17000', '20500'],
        ['10/08/2024', 'Despesas de Transporte', '', '1200', '20500', '19300']
    ]

    criar_tabela_formatada(doc, None, dados, adicionar_quebra=False)
    doc.save(caminho_completo)

    return f"[OK] Relatorio simples criado: {caminho_completo}"


def gerar_relatorio_completo(nome_arquivo="documento3.docx", caminho=None):
    """
    Gera relatorio completo com DUAS tabelas (20 registros em 2 paginas).
    Equivalente ao gera_documento3.py

    Args:
        nome_arquivo: Nome do arquivo a salvar
        caminho: Caminho para salvar (usa CAMINHO_SAIDA se None)
    """
    if caminho is None:
        caminho = CAMINHO_SAIDA

    caminho_completo = os.path.join(caminho, nome_arquivo)

    doc = Document()
    doc.add_heading('Relatorio Financeiro', level=1)

    # Dados do período 01-10 de agosto (Primeira tabela)
    dados_parte1 = [
        ['01/08/2025', 'Venda Produto A', '5000', '', '10000', '15000'],
        ['02/08/2025', 'Compra Insumos', '', '2000', '15000', '13000'],
        ['03/08/2025', 'Venda Produto B', '3000', '', '13000', '16000'],
        ['04/08/2025', 'Despesas Operacionais', '', '1500', '16000', '14500'],
        ['05/08/2025', 'Venda Produto C', '7000', '', '14500', '21500'],
        ['06/08/2025', 'Despesas de Marketing', '', '2500', '21500', '19000'],
        ['07/08/2025', 'Venda Produto D', '4000', '', '19000', '23000'],
        ['08/08/2025', 'Compra de Equipamentos', '', '6000', '23000', '17000'],
        ['09/08/2025', 'Venda Produto E', '3500', '', '17000', '20500'],
        ['10/08/2024', 'Despesas de Transporte', '', '1200', '20500', '19300']
    ]

    # Primeira tabela
    criar_tabela_formatada(doc, "Período: 01-10 de Agosto de 2025", dados_parte1, adicionar_quebra=True)

    # Dados do período 11-20 de agosto (Segunda tabela)
    dados_parte2 = [
        ['11/08/2025', 'Venda Produto F', '8000', '', '19300', '27300'],
        ['12/08/2025', 'Compra de Matéria-Prima', '', '4000', '27300', '23300'],
        ['13/08/2025', 'Venda Produto G', '4500', '', '23300', '27800'],
        ['14/08/2025', 'Despesas de Manutenção', '', '1700', '27800', '26100'],
        ['15/08/2025', 'Venda Produto H', '6000', '', '26100', '32100'],
        ['16/08/2025', 'Despesas Diversas', '', '2300', '32100', '29800'],
        ['17/08/2025', 'Venda Produto I', '7500', '', '29800', '37300'],
        ['18/08/2025', 'Compra de Ferramentas', '', '3200', '37300', '34100'],
        ['19/08/2025', 'Venda Produto J', '5000', '', '34100', '39100'],
        ['20/08/2025', 'Despesas Gerais', '', '1800', '39100', '37300']
    ]

    # Segunda tabela
    criar_tabela_formatada(doc, "Periodo: 11-20 de Agosto de 2025", dados_parte2, adicionar_quebra=False)

    doc.save(caminho_completo)

    return f"[OK] Relatorio completo criado: {caminho_completo}"


def gerar_relatorio_personalizado(nome_arquivo="documento_custom.docx", num_tabelas=2, caminho=None):
    """
    Gera relatorio personalizado com numero flexivel de tabelas.

    Args:
        nome_arquivo: Nome do arquivo a salvar
        num_tabelas: Numero de tabelas a gerar (1 ou 2)
        caminho: Caminho para salvar (usa CAMINHO_SAIDA se None)
    """
    if caminho is None:
        caminho = CAMINHO_SAIDA

    if num_tabelas == 1:
        return gerar_relatorio_simples(nome_arquivo, caminho)
    elif num_tabelas == 2:
        return gerar_relatorio_completo(nome_arquivo, caminho)
    else:
        return "Erro: num_tabelas deve ser 1 ou 2"


# ============================================================================
# EXECUÇÃO PRINCIPAL
# ============================================================================

if __name__ == "__main__":
    print("=" * 70)
    print("PROJETO 6 - GERACAO UNIFICADA DE RELATORIOS FINANCEIROS")
    print("=" * 70)
    print(f"\nCaminho de saida: {CAMINHO_SAIDA}\n")

    # Opcao 1: Gerar relatorio simples (1 tabela)
    print("[1] Gerando relatorio simples (1 tabela, 10 registros)...")
    resultado1 = gerar_relatorio_simples("documento2.docx", CAMINHO_SAIDA)
    print(resultado1)

    # Opcao 2: Gerar relatorio completo (2 tabelas)
    print("\n[2] Gerando relatorio completo (2 tabelas, 20 registros)...")
    resultado2 = gerar_relatorio_completo("documento3.docx", CAMINHO_SAIDA)
    print(resultado2)

    print("\n" + "=" * 70)
    print("RESUMO:")
    print("=" * 70)
    print(resultado1)
    print(resultado2)
    print("\nAmbos os relatorios foram gerados com sucesso!")
    print("=" * 70)
